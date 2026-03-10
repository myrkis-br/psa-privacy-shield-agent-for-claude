"""
PSA - Privacy Shield Agent
Script: anonymize_document.py
Responsável: PSA Guardião

Anonimiza documentos DOCX e TXT:
  - Detecta e substitui CPF, CNPJ, email, telefone, CEP, datas, valores monetários
  - Substitui nomes de pessoas e empresas por tokens [PESSOA_N] / [EMPRESA_N]
  - Extrai headers/footers de DOCX (H-06)
  - Fallback de encoding para TXT (H-06)
  - Seleciona amostra de parágrafos mais relevantes (conteúdo substancial)
  - Salva como TXT anonimizado em data/anonymized/
  - Salva mapa de entidades em data/maps/

Uso:
  python3 scripts/anonymize_document.py <caminho_do_arquivo> [--sample <n_paragrafos>]

Exemplos:
  python3 scripts/anonymize_document.py data/real/contrato.docx
  python3 scripts/anonymize_document.py data/real/relatorio.txt --sample 30
"""

import sys
import json
import argparse
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Tuple, Dict

# ---------------------------------------------------------------------------
# Configuração de caminhos
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent.parent
ANONYMIZED_DIR = BASE_DIR / "data" / "anonymized"
MAPS_DIR = BASE_DIR / "data" / "maps"
LOGS_DIR = BASE_DIR / "logs"

for d in (ANONYMIZED_DIR, MAPS_DIR, LOGS_DIR):
    d.mkdir(parents=True, exist_ok=True)

sys.path.insert(0, str(BASE_DIR / "scripts"))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [PSA-Guardião] %(levelname)s: %(message)s",
    handlers=[
        logging.FileHandler(LOGS_DIR / "psa_guardiao.log"),
        logging.StreamHandler(sys.stdout),
    ],
)
log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Imports de terceiros
# ---------------------------------------------------------------------------

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from text_engine import TextAnonymizer

# ---------------------------------------------------------------------------
# Extração de parágrafos
# ---------------------------------------------------------------------------

def _extract_paragraphs_docx(path: Path) -> List[Tuple[str, str]]:
    """
    Extrai parágrafos de um DOCX.
    H-06: Inclui headers, footers e text boxes.
    Retorna lista de (estilo, texto).
    """
    if DocxDocument is None:
        raise ImportError("python-docx não instalado. Execute: pip3 install python-docx")

    doc = DocxDocument(str(path))
    paragraphs = []

    # H-06: Extrai headers e footers
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header and header.paragraphs:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(("Header", text))
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer and footer.paragraphs:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(("Footer", text))

    # Parágrafos do corpo
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "Normal"
            paragraphs.append((style, text))

    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(("Table", row_text))

    return paragraphs


def _extract_paragraphs_txt(path: Path) -> List[Tuple[str, str]]:
    """
    Extrai parágrafos de um TXT.
    H-06: Fallback de encoding (utf-8 → latin-1 → cp1252).
    Parágrafos são blocos separados por linha em branco.
    """
    content = None
    for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1252", "iso-8859-1"):
        try:
            content = path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if content is None:
        content = path.read_text(encoding="utf-8", errors="replace")
        log.warning(f"Encoding não detectado para '{path.name}', usando UTF-8 com replace")

    blocks = []
    current = []

    for line in content.splitlines():
        stripped = line.strip()
        if stripped:
            current.append(stripped)
        elif current:
            blocks.append(("Normal", " ".join(current)))
            current = []

    if current:
        blocks.append(("Normal", " ".join(current)))

    return blocks


def _sample_paragraphs(
    paragraphs: List[Tuple[str, str]],
    max_paragraphs: int,
) -> List[Tuple[str, str]]:
    """
    Seleciona amostra representativa de parágrafos.

    Estratégia:
    1. Sempre inclui todos os headings (títulos de seção)
    2. Filtra parágrafos com conteúdo substancial (>30 chars)
    3. Distribui a amostra ao longo do documento (início, meio, fim)
    """
    heading_indices = {
        i for i, (s, _) in enumerate(paragraphs)
        if "heading" in s.lower() or "título" in s.lower()
    }
    content_indices = [
        i for i, (s, t) in enumerate(paragraphs)
        if len(t) >= 30 and i not in heading_indices
    ]

    if len(content_indices) <= max_paragraphs:
        return paragraphs  # documento curto, retorna tudo

    # Distribui: 40% início, 40% meio, 20% fim
    n = max(max_paragraphs - len(heading_indices), 5)

    n_start = int(n * 0.4)
    n_mid = int(n * 0.4)
    n_end = n - n_start - n_mid

    mid_start = len(content_indices) // 2 - n_mid // 2
    mid_end = mid_start + n_mid

    selected_indices = set(
        content_indices[:n_start]
        + content_indices[mid_start:mid_end]
        + content_indices[-n_end:]
    )
    selected_indices |= heading_indices

    result = [paragraphs[i] for i in sorted(selected_indices)]
    return result[:max_paragraphs + len(heading_indices)]


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def anonymize_document(
    input_path: Path,
    sample_paragraphs: int = 20,
) -> Tuple[Path, Path]:
    """
    Anonimiza um documento DOCX ou TXT.

    Args:
        input_path: Caminho para o arquivo em data/real/
        sample_paragraphs: Número máximo de parágrafos na amostra

    Returns:
        Tuple (caminho_anonimizado, caminho_mapa)
    """
    input_path = Path(input_path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = input_path.stem
    suffix = input_path.suffix.lower()

    log.info(f"Iniciando anonimização de documento: {input_path.name}")

    # --- Extração ---
    if suffix == ".docx":
        paragraphs = _extract_paragraphs_docx(input_path)
    elif suffix == ".txt":
        paragraphs = _extract_paragraphs_txt(input_path)
    else:
        raise ValueError(f"Formato não suportado: {suffix}. Use DOCX ou TXT.")

    total_paragraphs = len(paragraphs)
    log.info(f"Documento carregado: {total_paragraphs} parágrafos")

    # --- Amostragem ---
    sample = _sample_paragraphs(paragraphs, sample_paragraphs)
    log.info(f"Amostra selecionada: {len(sample)} de {total_paragraphs} parágrafos")

    # --- Anonimização ---
    engine = TextAnonymizer()
    anonymized_paragraphs = []

    for style, text in sample:
        anon_text = engine.anonymize(text)
        anonymized_paragraphs.append((style, anon_text))

    log.info(f"Entidades substituídas: {len(engine.entity_map)} ocorrências únicas")

    # M-08: Log sem PII — mostra apenas tokens, não originais
    for i, (token, _original) in enumerate(list(engine.entity_map.items())[:5]):
        log.info(f"  Substituição #{i+1}: -> '{token}'")

    # --- Montar texto de saída ---
    lines = []
    lines.append("=" * 70)
    lines.append("PSA - DOCUMENTO ANONIMIZADO")
    lines.append(f"Original: {input_path.name}")
    lines.append(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    lines.append(f"Parágrafos: {len(sample)} de {total_paragraphs} (amostra)")
    lines.append("=" * 70)
    lines.append("")

    for style, text in anonymized_paragraphs:
        if "heading" in style.lower():
            lines.append(f"\n## {text}\n")
        elif style == "Table":
            lines.append(f"[TABELA] {text}")
        elif style == "Header":
            lines.append(f"[CABEÇALHO] {text}")
        elif style == "Footer":
            lines.append(f"[RODAPÉ] {text}")
        else:
            lines.append(text)
            lines.append("")

    output_text = "\n".join(lines)

    # --- Salvar arquivo anonimizado ---
    anon_filename = f"anon_{stem}_{timestamp}.txt"
    anon_path = ANONYMIZED_DIR / anon_filename
    anon_path.write_text(output_text, encoding="utf-8")
    log.info(f"Documento anonimizado salvo: {anon_path}")

    # --- Salvar mapa ---
    map_data = {
        "timestamp": datetime.now().isoformat(),
        "tipo": "documento",
        "formato_original": suffix,
        "arquivo_original": str(input_path),
        "arquivo_anonimizado": str(anon_path),
        "total_paragrafos_original": total_paragraphs,
        "total_paragrafos_amostra": len(sample),
        "entidades": {
            token: original
            for token, original in engine.entity_map.items()
        },
        "estatisticas": {
            "pessoas_substituidas": engine._counters["pessoa"],
            "empresas_substituidas": engine._counters["empresa"],
            "total_entidades": len(engine.entity_map),
        },
    }

    map_filename = f"map_{stem}_{timestamp}.json"
    map_path = MAPS_DIR / map_filename
    with open(map_path, "w", encoding="utf-8") as f:
        json.dump(map_data, f, ensure_ascii=False, indent=2)

    log.info(f"Mapa salvo: {map_path}")
    log.info(
        f"Anonimização concluída: {engine._counters['pessoa']} pessoas, "
        f"{engine._counters['empresa']} empresas substituídas."
    )

    return anon_path, map_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="PSA Guardião — Anonimizador de documentos DOCX/TXT"
    )
    parser.add_argument("arquivo", help="Caminho para o arquivo a anonimizar")
    parser.add_argument(
        "--sample",
        type=int,
        default=20,
        metavar="N",
        help="Número máximo de parágrafos na amostra (padrão: 20)",
    )
    args = parser.parse_args()

    input_path = Path(args.arquivo)
    if not input_path.exists():
        log.error(f"Arquivo não encontrado: {input_path}")
        sys.exit(1)

    anon_path, map_path = anonymize_document(input_path, sample_paragraphs=args.sample)

    print("\n" + "=" * 60)
    print("PSA GUARDIÃO — ANONIMIZAÇÃO CONCLUÍDA")
    print("=" * 60)
    print(f"  Arquivo anonimizado    : {anon_path}")
    print(f"  Mapa de correspondência: {map_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
